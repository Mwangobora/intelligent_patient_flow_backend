from __future__ import annotations

import csv
from io import BytesIO, StringIO
from pathlib import Path

from django.conf import settings
from django.utils import timezone

from apps.reporting.models import ReportExport
from common.exceptions import NotFoundError, ValidationError


CONTENT_TYPES = {
    ReportExport.ExportFormat.CSV: "text/csv",
    ReportExport.ExportFormat.XLSX: "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ReportExport.ExportFormat.PDF: "application/pdf",
    ReportExport.ExportFormat.DOCX: "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def _private_storage_root() -> Path:
    root = getattr(settings, "PRIVATE_MEDIA_ROOT", settings.BASE_DIR / "private_media")
    return Path(root)


def _headers(rows: list[dict]) -> list[str]:
    if not rows:
        return ["message"]
    headers = []
    for row in rows:
        for key in row.keys():
            if key not in headers:
                headers.append(key)
    return headers


def _stringify(value):
    if value is None:
        return ""
    return str(value)


def _metadata_lines(*, title: str, metadata: dict) -> list[list[str]]:
    return [[title], ["generated_at", timezone.now().isoformat()]] + [[str(key), _stringify(value)] for key, value in metadata.items()]


def render_report_to_csv(*, title: str, rows: list[dict], metadata: dict) -> bytes:
    output = StringIO()
    writer = csv.writer(output)
    writer.writerows(_metadata_lines(title=title, metadata=metadata))
    writer.writerow([])
    headers = _headers(rows)
    writer.writerow(headers)
    if rows:
        for row in rows:
            writer.writerow([_stringify(row.get(header)) for header in headers])
    else:
        writer.writerow(["No data"])
    return output.getvalue().encode("utf-8")


def render_report_to_xlsx(*, title: str, rows: list[dict], metadata: dict) -> bytes:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font
    except ImportError as exc:
        raise ValidationError("XLSX generation dependency openpyxl is not installed.") from exc

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Report"
    sheet.append([title])
    sheet["A1"].font = Font(bold=True, size=14)
    for key, value in metadata.items():
        sheet.append([str(key), _stringify(value)])
    sheet.append(["generated_at", timezone.now().isoformat()])
    sheet.append([])

    headers = _headers(rows)
    sheet.append(headers)
    header_row = sheet.max_row
    for cell in sheet[header_row]:
        cell.font = Font(bold=True)
    if rows:
        for row in rows:
            sheet.append([_stringify(row.get(header)) for header in headers])
    else:
        sheet.append(["No data"])

    sheet.freeze_panes = f"A{header_row + 1}"
    for column_cells in sheet.columns:
        max_length = max(len(_stringify(cell.value)) for cell in column_cells)
        sheet.column_dimensions[column_cells[0].column_letter].width = min(max(max_length + 2, 12), 50)

    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def render_report_to_pdf(*, title: str, rows: list[dict], metadata: dict) -> bytes:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import landscape, letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError as exc:
        raise ValidationError("PDF generation dependency reportlab is not installed.") from exc

    output = BytesIO()
    document = SimpleDocTemplate(output, pagesize=landscape(letter), rightMargin=24, leftMargin=24, topMargin=24, bottomMargin=24)
    styles = getSampleStyleSheet()
    elements = [Paragraph(title, styles["Title"]), Spacer(1, 12)]
    for key, value in {**metadata, "generated_at": timezone.now().isoformat()}.items():
        elements.append(Paragraph(f"<b>{key}</b>: {_stringify(value)}", styles["Normal"]))
    elements.append(Spacer(1, 12))

    headers = _headers(rows)
    table_rows = [headers] + ([[row.get(header, "") for header in headers] for row in rows] if rows else [["No data"]])
    table = Table(table_rows, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9EAF7")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    elements.append(table)
    document.build(elements)
    return output.getvalue()


def render_report_to_docx(*, title: str, rows: list[dict], metadata: dict) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValidationError("DOCX generation dependency python-docx is not installed.") from exc

    document = Document()
    document.add_heading(title, level=1)
    for key, value in {**metadata, "generated_at": timezone.now().isoformat()}.items():
        document.add_paragraph(f"{key}: {_stringify(value)}")
    headers = _headers(rows)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    if rows:
        for row in rows:
            cells = table.add_row().cells
            for index, header in enumerate(headers):
                cells[index].text = _stringify(row.get(header))
    else:
        document.add_paragraph("No data")

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def render_report_file(*, report_export: ReportExport, title: str, rows: list[dict], metadata: dict) -> bytes:
    renderer = {
        ReportExport.ExportFormat.CSV: render_report_to_csv,
        ReportExport.ExportFormat.XLSX: render_report_to_xlsx,
        ReportExport.ExportFormat.PDF: render_report_to_pdf,
        ReportExport.ExportFormat.DOCX: render_report_to_docx,
    }[report_export.export_format]
    return renderer(title=title, rows=rows, metadata=metadata)


def save_report_file(*, report_export: ReportExport, content: bytes) -> str:
    storage_key = f"reports/{report_export.organization_id}/{report_export.id}.{report_export.export_format}"
    file_path = _private_storage_root() / storage_key
    file_path.parent.mkdir(parents=True, exist_ok=True)
    file_path.write_bytes(content)
    return storage_key


def get_report_file_path_or_content(*, storage_key: str) -> Path:
    file_path = _private_storage_root() / storage_key
    if not file_path.exists() or not file_path.is_file():
        raise NotFoundError("Report file not found.")
    return file_path
